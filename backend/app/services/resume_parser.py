from __future__ import annotations

import asyncio
import io
import os
import re
from dataclasses import asdict, dataclass, field
from typing import Any, Optional


_BULLET_LINE_RE = re.compile(r"^\s*(?:[\u2022\u00b7\u25cf\u25aa\u25e6*+-]|\d+[.)\u3001]|[\uff08(]?\d+[\uff09)])\s+")
_SECTION_HEADINGS = (
    "\u5de5\u4f5c\u7ecf\u5386",
    "\u5b9e\u4e60\u7ecf\u5386",
    "\u9879\u76ee\u7ecf\u5386",
    "\u6559\u80b2\u7ecf\u5386",
    "\u6821\u56ed\u7ecf\u5386",
    "\u4e2a\u4eba\u7ecf\u5386",
    "\u5b9e\u8df5\u7ecf\u5386",
    "\u83b7\u5956\u7ecf\u5386",
    "\u6280\u80fd",
    "\u8bc1\u4e66",
    "\u81ea\u6211\u8bc4\u4ef7",
    "\u4e2a\u4eba\u603b\u7ed3",
)
_SECTION_HEADING_RE = re.compile(
    r"^\s*(?:"
    + "|".join(re.escape(item) for item in _SECTION_HEADINGS)
    + r"|WORK\s+EXPERIENCE|EXPERIENCE|PROJECTS?|EDUCATION|SKILLS?|CERTIFICATIONS?"
    + r")\s*[:\uff1a]?\s*$",
    re.IGNORECASE,
)
_DATE_LINE_RE = re.compile(r"\b(?:19|20)\d{2}(?:[./-]\d{1,2})?\b")
_SPACING_RE = re.compile(r"[ \t\u00a0]+")
_MEANINGFUL_CHAR_RE = re.compile(r"[\u3400-\u9fffA-Za-z0-9]")
_REPEATED_MARGIN_MIN_PAGES = 2
_OCR_MIN_CHAR_COUNT = 80
_OCR_MIN_QUALITY = 0.42
_LIGATURES = str.maketrans(
    {
        "\ufb00": "ff",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
    }
)


@dataclass
class ResumeParsePage:
    page_number: int
    text: str
    method: str
    char_count: int
    quality_score: float
    warnings: list[str] = field(default_factory=list)

    def public_dict(self) -> dict[str, Any]:
        payload = asdict(self)
        payload.pop("text", None)
        return payload


@dataclass
class ResumeParseResult:
    text: str
    parser: str
    pages: list[ResumeParsePage]
    warnings: list[str] = field(default_factory=list)

    @property
    def used_ocr(self) -> bool:
        return any(page.method == "ocr" for page in self.pages)

    @property
    def average_quality(self) -> float:
        if not self.pages:
            return 0.0
        return round(sum(page.quality_score for page in self.pages) / len(self.pages), 3)

    def public_dict(self) -> dict[str, Any]:
        return {
            "parser": self.parser,
            "page_count": len(self.pages),
            "used_ocr": self.used_ocr,
            "average_quality": self.average_quality,
            "low_quality_pages": [
                page.page_number
                for page in self.pages
                if page.quality_score < _OCR_MIN_QUALITY
            ],
            "warnings": self.warnings,
            "pages": [page.public_dict() for page in self.pages],
        }


@dataclass
class _PdfTextBlock:
    text: str
    bbox: tuple[float, float, float, float]


def _clean_pdf_text(text: str) -> str:
    text = (
        (text or "")
        .translate(_LIGATURES)
        .replace("\u00ad", "")
        .replace("\u00a0", " ")
        .replace("\u200b", "")
    )
    text = _SPACING_RE.sub(" ", text)
    return text.strip()


def _looks_like_hard_break(line: str) -> bool:
    if not line:
        return True
    if _BULLET_LINE_RE.match(line) or _SECTION_HEADING_RE.match(line):
        return True
    if _DATE_LINE_RE.search(line):
        return True
    if re.search(r"(?:\u81f3\u4eca|Present|present)\s*$", line):
        return True
    return False


def _is_sentence_end(text: str) -> bool:
    return text.endswith(("\u3002", "\uff01", "\uff1f", ".", "!", "?", ";", "\uff1b", ":", "\uff1a"))


def _join_visual_wrap(left: str, right: str) -> str:
    if re.search(r"[A-Za-z]-$", left) and re.match(r"^[a-z]", right):
        return f"{left[:-1]}{right}"
    return f"{left} {right}"


def _normalize_extracted_text(text: str) -> str:
    """Join PDF visual wraps while keeping real headings and bullet items."""
    lines = [_clean_pdf_text(line) for line in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    blocks: list[str] = []
    current = ""

    for line in lines:
        if not line:
            if current:
                blocks.append(current.strip())
                current = ""
            continue

        if _looks_like_hard_break(line):
            if current:
                blocks.append(current.strip())
            current = line
            continue

        if not current:
            current = line
            continue

        if _SECTION_HEADING_RE.match(current) or _is_sentence_end(current):
            blocks.append(current.strip())
            current = line
        else:
            current = _join_visual_wrap(current, line)

    if current:
        blocks.append(current.strip())

    return "\n".join(blocks)


def _block_text(block: Any) -> str:
    if isinstance(block, dict):
        lines: list[str] = []
        for line in block.get("lines") or []:
            parts = [_clean_pdf_text(str(span.get("text") or "")) for span in line.get("spans") or []]
            text = _clean_pdf_text("".join(parts))
            if text:
                lines.append(text)
        return "\n".join(lines)
    if isinstance(block, (tuple, list)) and len(block) >= 5:
        return str(block[4] or "")
    return ""


def _block_bbox(block: Any) -> tuple[float, float, float, float]:
    if isinstance(block, dict):
        bbox = block.get("bbox") or (0, 0, 0, 0)
    elif isinstance(block, (tuple, list)) and len(block) >= 4:
        bbox = block[:4]
    else:
        bbox = (0, 0, 0, 0)
    try:
        return float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
    except Exception:
        return 0.0, 0.0, 0.0, 0.0


def _sort_pdf_blocks(blocks: list[Any], page_width: float = 0.0) -> list[Any]:
    """Keep a detected two-column resume column-stable instead of interleaving rows."""
    if len(blocks) < 4 or page_width <= 0:
        return sorted(blocks, key=lambda block: (_block_bbox(block)[1], _block_bbox(block)[0]))

    valid = [block for block in blocks if _block_bbox(block)[2] > _block_bbox(block)[0]]
    if len(valid) < 4:
        return sorted(blocks, key=lambda block: (_block_bbox(block)[1], _block_bbox(block)[0]))

    midpoint = page_width / 2
    tolerance = max(8.0, page_width * 0.018)
    left = [block for block in valid if _block_bbox(block)[2] < midpoint + tolerance]
    right = [block for block in valid if _block_bbox(block)[0] > midpoint - tolerance]
    spanning = [block for block in valid if block not in left and block not in right]
    if len(left) < 2 or len(right) < 2:
        return sorted(blocks, key=lambda block: (_block_bbox(block)[1], _block_bbox(block)[0]))

    left_top = min(_block_bbox(block)[1] for block in left)
    right_top = min(_block_bbox(block)[1] for block in right)
    left_bottom = max(_block_bbox(block)[3] for block in left)
    right_bottom = max(_block_bbox(block)[3] for block in right)
    overlap = max(0.0, min(left_bottom, right_bottom) - max(left_top, right_top))
    shorter_height = max(1.0, min(left_bottom - left_top, right_bottom - right_top))
    if overlap / shorter_height < 0.2:
        return sorted(blocks, key=lambda block: (_block_bbox(block)[1], _block_bbox(block)[0]))

    column_top = min(left_top, right_top)
    column_bottom = max(left_bottom, right_bottom)
    top_spanning = [block for block in spanning if _block_bbox(block)[3] <= column_top + tolerance]
    bottom_spanning = [block for block in spanning if _block_bbox(block)[1] >= column_bottom - tolerance]
    middle_spanning = [
        block
        for block in spanning
        if block not in top_spanning and block not in bottom_spanning
    ]
    by_position = lambda block: (_block_bbox(block)[1], _block_bbox(block)[0])
    return (
        sorted(top_spanning, key=by_position)
        + sorted(left, key=by_position)
        + sorted(right, key=by_position)
        + sorted(middle_spanning + bottom_spanning, key=by_position)
    )


def _text_quality(text: str) -> tuple[int, float]:
    compact = re.sub(r"\s+", "", text or "")
    char_count = len(compact)
    if char_count == 0:
        return 0, 0.0
    meaningful_ratio = len(_MEANINGFUL_CHAR_RE.findall(compact)) / char_count
    replacement_ratio = compact.count("\ufffd") / char_count
    length_score = min(1.0, char_count / 220)
    line_score = min(1.0, len([line for line in (text or "").splitlines() if line.strip()]) / 8)
    score = length_score * 0.5 + meaningful_ratio * 0.4 + line_score * 0.1 - replacement_ratio * 2
    return char_count, round(max(0.0, min(1.0, score)), 3)


def _page_needs_ocr(text: str) -> bool:
    char_count, quality_score = _text_quality(text)
    return char_count < _OCR_MIN_CHAR_COUNT or quality_score < _OCR_MIN_QUALITY


def _extract_page_blocks(page: Any, textpage: Any = None) -> list[_PdfTextBlock]:
    try:
        kwargs = {"sort": False}
        if textpage is not None:
            kwargs["textpage"] = textpage
        raw_blocks = page.get_text("dict", **kwargs).get("blocks", [])
        text_blocks = [block for block in raw_blocks if block.get("type") == 0]
    except Exception:
        kwargs = {"sort": False}
        if textpage is not None:
            kwargs["textpage"] = textpage
        text_blocks = page.get_text("blocks", **kwargs)

    result: list[_PdfTextBlock] = []
    for block in _sort_pdf_blocks(text_blocks, float(page.rect.width)):
        text = _normalize_extracted_text(_block_text(block))
        if text:
            result.append(_PdfTextBlock(text=text, bbox=_block_bbox(block)))
    return result


def _margin_key(text: str) -> str:
    return re.sub(r"\d+", "#", re.sub(r"\s+", " ", text.strip().lower()))


def _remove_repeated_margins(
    pages: list[tuple[list[_PdfTextBlock], float]],
) -> list[list[_PdfTextBlock]]:
    if len(pages) < _REPEATED_MARGIN_MIN_PAGES:
        return [blocks for blocks, _ in pages]

    counts: dict[str, int] = {}
    for blocks, page_height in pages:
        page_keys = {
            _margin_key(block.text)
            for block in blocks
            if block.text
            and (
                block.bbox[1] <= page_height * 0.08
                or block.bbox[3] >= page_height * 0.92
            )
        }
        for key in page_keys:
            counts[key] = counts.get(key, 0) + 1

    threshold = max(_REPEATED_MARGIN_MIN_PAGES, (len(pages) + 1) // 2)
    repeated = {key for key, count in counts.items() if count >= threshold}
    if not repeated:
        return [blocks for blocks, _ in pages]

    return [
        [
            block
            for block in blocks
            if not (
                _margin_key(block.text) in repeated
                and (
                    block.bbox[1] <= page_height * 0.08
                    or block.bbox[3] >= page_height * 0.92
                )
            )
        ]
        for blocks, page_height in pages
    ]


def _ocr_page_blocks(page: Any) -> tuple[list[_PdfTextBlock], str]:
    preferred = os.getenv("OFFERU_OCR_LANGUAGE", "chi_sim+eng").strip() or "chi_sim+eng"
    languages = [preferred]
    if preferred != "eng":
        languages.append("eng")
    last_error = ""
    for language in languages:
        try:
            textpage = page.get_textpage_ocr(language=language, dpi=200, full=True)
            return _extract_page_blocks(page, textpage=textpage), language
        except Exception as exc:
            last_error = str(exc)
    raise RuntimeError(last_error or "OCR unavailable")


def _parse_pdf_with_pymupdf(file_bytes: bytes) -> ResumeParseResult:
    import fitz

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    native_pages: list[tuple[list[_PdfTextBlock], float]] = []
    try:
        for page in doc:
            native_pages.append((_extract_page_blocks(page), float(page.rect.height)))
        cleaned_native_pages = _remove_repeated_margins(native_pages)

        pages: list[ResumeParsePage] = []
        warnings: list[str] = []
        for page_index, page in enumerate(doc):
            native_blocks = cleaned_native_pages[page_index]
            native_text = "\n".join(block.text for block in native_blocks if block.text)
            method = "native"
            page_warnings: list[str] = []
            text = native_text

            if _page_needs_ocr(native_text):
                try:
                    ocr_blocks, language = _ocr_page_blocks(page)
                    ocr_text = "\n".join(block.text for block in ocr_blocks if block.text)
                    if _text_quality(ocr_text)[1] > _text_quality(native_text)[1]:
                        text = ocr_text
                        method = "ocr"
                        page_warnings.append(f"page_ocr:{language}")
                    else:
                        page_warnings.append("ocr_not_better")
                except Exception:
                    page_warnings.append("ocr_unavailable")

            char_count, quality_score = _text_quality(text)
            if quality_score < _OCR_MIN_QUALITY:
                page_warnings.append("low_text_quality")
                warnings.append(f"page_{page_index + 1}_low_text_quality")
            if "ocr_unavailable" in page_warnings:
                warnings.append(f"page_{page_index + 1}_ocr_unavailable")

            pages.append(
                ResumeParsePage(
                    page_number=page_index + 1,
                    text=text,
                    method=method,
                    char_count=char_count,
                    quality_score=quality_score,
                    warnings=page_warnings,
                )
            )
    finally:
        doc.close()

    return ResumeParseResult(
        text="\n\n".join(page.text for page in pages if page.text.strip()),
        parser="pymupdf",
        pages=pages,
        warnings=list(dict.fromkeys(warnings)),
    )


def _parse_pdf_with_pypdf(file_bytes: bytes) -> ResumeParseResult:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[ResumeParsePage] = []
    for page_index, page in enumerate(reader.pages):
        text = _normalize_extracted_text(page.extract_text() or "")
        char_count, quality_score = _text_quality(text)
        warnings = ["pypdf_fallback"]
        if quality_score < _OCR_MIN_QUALITY:
            warnings.append("low_text_quality")
        pages.append(
            ResumeParsePage(
                page_number=page_index + 1,
                text=text,
                method="pypdf",
                char_count=char_count,
                quality_score=quality_score,
                warnings=warnings,
            )
        )
    return ResumeParseResult(
        text="\n\n".join(page.text for page in pages if page.text.strip()),
        parser="pypdf",
        pages=pages,
        warnings=["pymupdf_failed", "pypdf_fallback"],
    )


def _parse_pdf_document_sync(file_bytes: bytes) -> ResumeParseResult:
    try:
        result = _parse_pdf_with_pymupdf(file_bytes)
    except Exception:
        return _parse_pdf_with_pypdf(file_bytes)
    if result.text.strip():
        return result
    try:
        fallback = _parse_pdf_with_pypdf(file_bytes)
    except Exception:
        return result
    if fallback.text.strip():
        return fallback
    return result


def _extract_pdf_with_pymupdf(file_bytes: bytes) -> str:
    return _parse_pdf_with_pymupdf(file_bytes).text


def _extract_pdf_with_pypdf(file_bytes: bytes) -> str:
    return _parse_pdf_with_pypdf(file_bytes).text


def _parse_pdf_sync(file_bytes: bytes) -> str:
    """Backward-compatible plain-text PDF parser."""
    return _parse_pdf_document_sync(file_bytes).text


def _parse_docx_sync(file_bytes: bytes) -> str:
    """Extract plain text from Word (.docx)."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    texts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))
    return "\n".join(texts)


def _parse_docx_document_sync(file_bytes: bytes) -> ResumeParseResult:
    text = _parse_docx_sync(file_bytes)
    char_count, quality_score = _text_quality(text)
    warnings = ["low_text_quality"] if quality_score < _OCR_MIN_QUALITY else []
    return ResumeParseResult(
        text=text,
        parser="python-docx",
        pages=[
            ResumeParsePage(
                page_number=1,
                text=text,
                method="docx",
                char_count=char_count,
                quality_score=quality_score,
                warnings=warnings,
            )
        ],
        warnings=warnings,
    )


async def parse_resume_document(filename: str, file_bytes: bytes) -> Optional[ResumeParseResult]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return await asyncio.to_thread(_parse_pdf_document_sync, file_bytes)
    if lower.endswith(".docx"):
        return await asyncio.to_thread(_parse_docx_document_sync, file_bytes)
    return None


def locate_resume_source_pages(
    result: ResumeParseResult,
    payload: Any,
) -> list[int]:
    """Locate exact candidate fragments in page text without inventing provenance."""
    fragments: list[str] = []

    def collect(value: Any) -> None:
        if isinstance(value, str):
            text = _clean_pdf_text(value)
            if len(re.sub(r"\s+", "", text)) >= 6:
                fragments.append(text)
            return
        if isinstance(value, dict):
            for item in value.values():
                collect(item)
            return
        if isinstance(value, list):
            for item in value:
                collect(item)

    collect(payload)
    probes = sorted(
        {
            re.sub(r"[\W_]+", "", fragment.lower())
            for fragment in fragments
            if fragment
        },
        key=len,
        reverse=True,
    )[:12]
    if not probes:
        return []

    matched: list[int] = []
    for page in result.pages:
        page_text = re.sub(r"[\W_]+", "", page.text.lower())
        if any(len(probe) >= 6 and probe in page_text for probe in probes):
            matched.append(page.page_number)
    return matched


async def parse_resume_file(filename: str, file_bytes: bytes) -> Optional[str]:
    result = await parse_resume_document(filename, file_bytes)
    return result.text if result else None
